import json
import os
from typing import Optional

import docx
from docxcompose.composer import Composer

from utilities.patterns import Pattern, pattern_list
from db.storage import get_scenario
from db.connection import get_session


db_session = get_session()


class Scenario:
    def __init__(self, scenario=None, user='') -> None:
        self.data = list()
        self.name = ''
        self.user = user
        if scenario is None:
            for pattern in pattern_list:
                new = dict()
                new["name"] = pattern.name
                new["p_id"] = pattern.pattern
                new["enabled"] = False
                new["components"] = list()
                for sub in pattern.data["components"]:
                    new["components"].append(dict())
                    new["components"][-1]["ps_id"] = sub["id"]
                    new["components"][-1]["name"] = sub["name"]
                    new["components"][-1]["enabled"] = False
                    new["components"][-1]["role"] = "Любой пользователь"
                    new["components"][-1]["components"] = list()
                    for sub2 in sub["components"]:
                        new["components"][-1]["components"].append(dict())
                        new["components"][-1]["components"][-1]["ps2_id"] = sub2["id"]
                        new["components"][-1]["components"][-1]["name"] = sub2["name"]
                        new["components"][-1]["components"][-1]["enabled"] = 0
                self.data.append(new)
        else:
            scenario = get_scenario(scenario, db_session)
            self.name = scenario["name"]
            self.id = scenario["id"]
            self.data = scenario["components"]

    def build_docx(self):
        dct = dict()
        for i in self.data:
            dct[i["p_id"]] = dict()
            dct[i["p_id"]]['enabled'] = i["enabled"]
            dct[i["p_id"]]['components'] = dict()
            for j in i["components"]:
                cur = dct[i["p_id"]]["components"]
                cur[j["ps_id"]] = dict()
                cur[j["ps_id"]]['enabled'] = j["enabled"]
                cur[j["ps_id"]]["components"] = dict()
                cur[j["ps_id"]]["role"] = j["role"]
                for k in j["components"]:
                    cur2 = cur[j["ps_id"]]["components"]
                    cur2[k["ps2_id"]] = dict()
                    cur2[k["ps2_id"]]['enabled'] = k["enabled"]
        doc = docx.Document()
        composer = Composer(doc)
        doc.add_heading(self.name, 0)
        for pattern in pattern_list:
            if dct[pattern.pattern]["enabled"]:
                doc.add_heading(pattern.name, 1)
                for sub in pattern.data["components"]:
                    if dct[pattern.pattern]["components"][sub["id"]]["enabled"]:
                        doc.add_heading(sub["name"], 2)
                        for sub2 in sub["components"]:
                            if dct[pattern.pattern]["components"][sub["id"]]["components"][sub2["id"]]["enabled"]:
                                sub2["doc"].paragraphs[0].text = sub2["doc"].paragraphs[0].text.replace('пользователь',
                                                                dct[pattern.pattern]["components"][sub["id"]]["role"]
                                                                                                        )
                                doc.add_heading(
                                    sub2["name"], 3
                                )

                                composer.append(sub2["doc"])
        doc.save(f'data_files/user_data/{self.user}/{self.id}.docx')
